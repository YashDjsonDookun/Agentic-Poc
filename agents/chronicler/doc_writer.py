"""
Doc Writer (4.2): generate runbook/SOP from a cluster of closed incidents.
Outputs three formats: .md, .docx, .pdf — all saved to knowledge/generated/.

LLM placeholder: if rag.yaml has an LLM endpoint configured, call it; otherwise use templates.
"""
from __future__ import annotations

import textwrap
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from shared.config_loader import PROJECT_ROOT, get_rag_config

GENERATED = PROJECT_ROOT / "knowledge" / "generated"


def _ensure_dir() -> None:
    GENERATED.mkdir(parents=True, exist_ok=True)


def _build_md(cluster: dict, trace_rows: list[dict]) -> str:
    """Build markdown content from a cluster of incidents + their trace data."""
    service = cluster.get("service", "unknown")
    theme = cluster.get("theme", "general")
    count = cluster.get("count", 0)
    severities = ", ".join(cluster.get("severities", []))
    summaries = cluster.get("summaries", [])
    incidents = cluster.get("incidents", [])

    rca_hypotheses: list[str] = []
    resolution_steps: list[str] = []
    recommendations: list[str] = []

    for row in trace_rows:
        agent = (row.get("agent") or "").lower()
        action = (row.get("action") or "").lower()
        rationale = row.get("rationale", "")
        detail = row.get("detail", "")

        if "rca" in agent:
            if rationale:
                rca_hypotheses.append(rationale)
        if "recommender" in agent or "recommend" in action:
            if rationale:
                recommendations.append(rationale)
            if detail:
                recommendations.append(detail)
        if "executor" in agent or "execute" in action:
            if rationale:
                resolution_steps.append(rationale)
        if "closer" in agent:
            if rationale:
                resolution_steps.append(rationale)

    rca_hypotheses = list(dict.fromkeys(rca_hypotheses))
    resolution_steps = list(dict.fromkeys(resolution_steps))
    recommendations = list(dict.fromkeys(recommendations))

    title = f"{service.replace('-', ' ').title()} - {theme.replace('_', ' ').title()}"
    now = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")

    lines = [
        f"# {title}",
        "",
        f"> Auto-generated on {now} from {count} closed incident(s).",
        "",
        "## Overview",
        "",
        f"- **Service:** {service}",
        f"- **Theme:** {theme.replace('_', ' ')}",
        f"- **Severity levels seen:** {severities or 'N/A'}",
        f"- **Incident count:** {count}",
        "",
        "## Symptoms",
        "",
    ]
    for s in summaries:
        lines.append(f"- {s}")
    if not summaries:
        lines.append("- _(no summaries available)_")
    lines += ["", "## Root Cause Analysis", ""]
    if rca_hypotheses:
        for h in rca_hypotheses:
            lines.append(f"- {h}")
    else:
        lines.append("- No RCA data available — consider manual investigation.")
    lines += ["", "## Resolution Steps", ""]
    if resolution_steps:
        for idx, step in enumerate(resolution_steps, 1):
            lines.append(f"{idx}. {step}")
    else:
        lines.append("1. Review incident details and apply appropriate runbook.")
        lines.append("2. Verify service health post-resolution.")
    lines += ["", "## Rollback / Mitigation", ""]
    lines.append("- If the resolution fails, revert any changes and escalate.")
    lines.append("- Monitor affected service for 30 minutes post-resolution.")
    lines += ["", "## Prevention / Recommendations", ""]
    if recommendations:
        for r in recommendations:
            lines.append(f"- {r}")
    else:
        lines.append("- Review threshold configurations and alert rules.")
        lines.append("- Ensure runbooks are kept up to date with recent incident learnings.")

    lines += [
        "",
        "---",
        "",
        "## Incident Details",
        "",
        "| Incident ID | Ticket | Service | Severity | Summary |",
        "|------------|--------|---------|----------|---------|",
    ]
    for inc in incidents:
        iid = inc.get("incident_id", "")
        tn = inc.get("ticket_number", "") or inc.get("ticket_id", "")
        svc = inc.get("service", "")
        sev = inc.get("severity", "")
        summ = inc.get("summary", "")
        lines.append(f"| {iid} | {tn} | {svc} | {sev} | {summ} |")

    lines.append("")
    return "\n".join(lines)


def _write_md(content: str, path: Path) -> None:
    path.write_text(content, encoding="utf-8")


def _write_docx(content: str, path: Path) -> None:
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return

    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    for line in content.split("\n"):
        stripped = line.strip()
        if stripped.startswith("# "):
            p = doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            p = doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("> "):
            p = doc.add_paragraph(stripped[2:])
            p.style = doc.styles["Intense Quote"] if "Intense Quote" in doc.styles else doc.styles["Normal"]
        elif stripped.startswith("- **"):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped.startswith("- "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped and stripped[0].isdigit() and ". " in stripped:
            text_part = stripped.split(". ", 1)[1] if ". " in stripped else stripped
            doc.add_paragraph(text_part, style="List Number")
        elif stripped.startswith("|") and "---" not in stripped:
            cells = [c.strip() for c in stripped.split("|")[1:-1]]
            if cells:
                doc.add_paragraph(" | ".join(cells))
        elif stripped == "---":
            doc.add_paragraph("─" * 50)
        elif stripped:
            doc.add_paragraph(stripped)

    doc.save(str(path))


def _sanitize_for_pdf(text: str) -> str:
    """Replace Unicode chars unsupported by built-in PDF fonts with ASCII equivalents."""
    replacements = {
        "\u2014": "--",   # em-dash
        "\u2013": "-",    # en-dash
        "\u2018": "'",    # left single quote
        "\u2019": "'",    # right single quote
        "\u201c": '"',    # left double quote
        "\u201d": '"',    # right double quote
        "\u2026": "...",  # ellipsis
        "\u2022": "*",    # bullet
        "\u2502": "|",    # box drawing vertical
        "\u2500": "-",    # box drawing horizontal
        "\u25cf": "*",    # black circle
        "\u00b7": ".",    # middle dot
    }
    for char, repl in replacements.items():
        text = text.replace(char, repl)
    return text.encode("latin-1", errors="replace").decode("latin-1")


def _write_pdf(content: str, path: Path) -> None:
    try:
        from fpdf import FPDF
    except ImportError:
        return

    content = _sanitize_for_pdf(content)

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Helvetica", size=10)

    full_w = pdf.w - pdf.l_margin - pdf.r_margin

    for line in content.split("\n"):
        stripped = line.strip()
        if not stripped:
            pdf.ln(3)
            continue
        if stripped.startswith("# "):
            pdf.set_x(pdf.l_margin)
            pdf.set_font("Helvetica", "B", 16)
            pdf.cell(full_w, 10, stripped[2:], new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("Helvetica", size=10)
        elif stripped.startswith("## "):
            pdf.set_x(pdf.l_margin)
            pdf.set_font("Helvetica", "B", 13)
            pdf.cell(full_w, 8, stripped[3:], new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("Helvetica", size=10)
        elif stripped.startswith("> "):
            pdf.set_x(pdf.l_margin)
            pdf.set_font("Helvetica", "I", 9)
            pdf.multi_cell(full_w, 5, stripped[2:])
            pdf.set_font("Helvetica", size=10)
        elif stripped.startswith("- ") or (stripped and stripped[0].isdigit() and ". " in stripped):
            indent = 5
            pdf.set_x(pdf.l_margin + indent)
            pdf.multi_cell(full_w - indent, 5, stripped)
        elif stripped.startswith("|") and "---" not in stripped:
            cells = [c.strip() for c in stripped.split("|")[1:-1]]
            pdf.set_x(pdf.l_margin)
            pdf.set_font("Helvetica", size=8)
            pdf.cell(full_w, 5, "  |  ".join(cells), new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("Helvetica", size=10)
        elif stripped == "---":
            pdf.set_x(pdf.l_margin)
            pdf.cell(full_w, 5, "_" * 80, new_x="LMARGIN", new_y="NEXT")
        else:
            pdf.set_x(pdf.l_margin)
            pdf.multi_cell(full_w, 5, stripped)

    pdf.output(str(path))


def _build_filename(cluster: dict) -> str:
    """Build a descriptive filename from cluster data.
    Format: {ticket_number}_{theme}_{severity}_{YYYY-MM-DD_HHmm}
    Examples:
      INC0010042_high_cpu_critical_2026-02-24_1430
      INC0010042_+3_error_spike_high_2026-02-24_1430  (when multiple incidents)
    """
    incidents = cluster.get("incidents", [])
    if not incidents:
        return f"{cluster.get('service', 'unknown')}_{cluster.get('theme', 'general')}"

    sorted_incs = sorted(incidents, key=lambda i: i.get("timestamp", ""), reverse=True)
    latest = sorted_incs[0]

    ticket = (latest.get("ticket_number") or "").strip()
    if not ticket:
        ticket = (latest.get("incident_id") or "unknown").strip()

    theme = cluster.get("theme", "general")

    top_sev = "medium"
    sev_rank = {"critical": 0, "high": 1, "medium": 2, "low": 3}
    for inc in incidents:
        s = (inc.get("severity") or "medium").lower()
        if sev_rank.get(s, 9) < sev_rank.get(top_sev, 9):
            top_sev = s

    try:
        from datetime import datetime as _dt
        ts = _dt.fromisoformat(latest.get("timestamp", "").replace("Z", "+00:00"))
        ts_str = ts.strftime("%Y-%m-%d_%H%M")
    except (ValueError, TypeError):
        ts_str = "undated"

    count = len(incidents)
    count_tag = f"_+{count - 1}" if count > 1 else ""

    parts = [ticket, count_tag.lstrip("_") if count_tag else None, theme, top_sev, ts_str]
    name = "_".join(p for p in parts if p)
    name = name.replace(" ", "_").replace("/", "_").replace("\\", "_").replace(":", "")
    return name


def generate_docs(
    cluster: dict,
    trace_rows: list[dict],
    name_override: str = "",
) -> dict[str, Path]:
    """Generate .md, .docx, and .pdf for a cluster. Returns {format: path}."""
    _ensure_dir()

    basename = name_override or _build_filename(cluster)
    basename = basename.replace(" ", "_").replace("/", "_")

    md_content = _build_md(cluster, trace_rows)

    paths: dict[str, Path] = {}

    md_path = GENERATED / f"{basename}.md"
    _write_md(md_content, md_path)
    paths["md"] = md_path

    docx_path = GENERATED / f"{basename}.docx"
    _write_docx(md_content, docx_path)
    if docx_path.exists():
        paths["docx"] = docx_path

    pdf_path = GENERATED / f"{basename}.pdf"
    _write_pdf(md_content, pdf_path)
    if pdf_path.exists():
        paths["pdf"] = pdf_path

    return paths
